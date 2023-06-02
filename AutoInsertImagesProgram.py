import tkinter as tk
from tkinterdnd2 import DND_FILES, TkinterDnD
from docx import Document
from collections import Counter
import requests
import os
import cv2
import numpy as np
from docx.shared import Inches
import io
import logging
from concurrent.futures import ThreadPoolExecutor
from cachetools import cached, TTLCache  # for caching
from tqdm import tqdm  # for progress bar
import asyncio  # for concurrency
import unittest  # for unit testing
from wikimedia import Wikimedia
import json
import tkinter.messagebox as messagebox
import re

# Set up logging
logging.basicConfig(level=logging.INFO)

# Load configuration
with open('config.json') as f:
    config = json.load(f)

# Create a ThreadPoolExecutor for fetching images in parallel
executor = ThreadPoolExecutor(max_workers=config['max_workers'])

# Set up caching
cache = TTLCache(maxsize=config['cache_size'], ttl=300)

# Create a Wikimedia object for interacting with the Wikimedia API
wikimedia = Wikimedia()

# List to store image URLs for bibliography
bibliography = []

@cached(cache)
def fetch_image(image_title):
    """
    Function to fetch the image URL
    """
    image_url = wikimedia.image_url(image_title)
    return image_url

def insert_image(doc, image_url, image_size):
    """
    Function to insert the image into the document
    """
    try:
        image_response = requests.get(image_url)
        image_bytes = image_response.content
        nparr = np.fromstring(image_bytes, np.uint8)
        img_np = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        resized_image = cv2.resize(img_np, (Inches(image_size), Inches(image_size)), interpolation = cv2.INTER_AREA)
        temp_filepath = 'temp.jpg'
        cv2.imwrite(temp_filepath, resized_image)
        doc.add_picture(temp_filepath)
        os.remove(temp_filepath)
        # Add image caption and source
        doc.add_paragraph(f"Image source: {image_url}")
        # Add image URL to bibliography
        bibliography.append(image_url)
    except requests.exceptions.RequestException as e:
        logging.error(f"An error occurred while fetching the image: {e}")
    except Exception as e:
        logging.error(f"An error occurred while inserting an image: {e}")

async def process_files(filepaths, image_size):
    """
    Function to process the files
    """
    for filepath in filepaths:
        try:
            if not os.path.isfile(filepath):
                raise FileNotFoundError(f"No such file or directory: '{filepath}'")
            if not re.search(r'\.docx$', filepath, re.I):
                raise ValueError(f"File '{filepath}' is not a Word document.")

            doc = Document(filepath)
            text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
            keywords = extract_keywords(text)
            for keyword in tqdm(keywords, desc="Processing keywords"):  # progress bar
                images = find_images(keyword)
                image_urls = list(executor.map(fetch_image, images))
                for image_url in image_urls:
                    insert_image(doc, image_url, image_size)
            # Add bibliography to the end of the document
            doc.add_section()
            doc.add_paragraph("Image Biography:")
            for source in bibliography:
                doc.add_paragraph(source)
            modified_filepath = 'modified_' + filepath
            doc.save(modified_filepath)
            logging.info(f"Finished processing file: {modified_filepath}")
        except Exception as e:
            logging.error(f"An error occurred while processing the file: {e}")
            messagebox.showerror("Error", f"An error occurred while processing the file: {e}")

def extract_keywords(text, threshold=5):
    """
    Function to extract keywords from the text
    """
    words = text.split()
    counter = Counter(words)
    keywords = [word for word,count in counter.items() if count > threshold]
    return keywords

def find_images(keyword):
    """
    Function to find images related to the keyword
    """
    images = wikimedia.search_images(keyword)
    return images

def drop(event):
    """
    Function to handle the drop event
    """
    try:
        filepaths = event.data.split()
        asyncio.run(process_files(filepaths, image_size=float(image_size_entry.get())))
    except Exception as e:
        logging.error(f"An error occurred while handling the drop event: {e}")
        messagebox.showerror("Error", f"An error occurred while processing the files: {e}")

# Add a simple GUI for keyword extraction
def create_gui():
    root = TkinterDnD.Tk()
    root.title("Document Processor")

    threshold_label = tk.Label(root, text="Keyword Threshold:")
    threshold_label.pack()
    threshold_entry = tk.Entry(root)
    threshold_entry.pack()
    threshold_button = tk.Button(root, text="Set Threshold", command=lambda: extract_keywords(threshold=int(threshold_entry.get())))
    threshold_button.pack()

    image_size_label = tk.Label(root, text="Image Size (in inches):")
    image_size_label.pack()
    global image_size_entry
    image_size_entry = tk.Entry(root)
    image_size_entry.pack()

    drop_label = tk.Label(root, text="Drop Files Here:")
    drop_label.pack()
    drop_frame = tk.Frame(root, height=100, width=100, bd=1, relief='sunken')
    drop_frame.drop_target_register(DND_FILES)
    drop_frame.dnd_bind('<<Drop>>', drop)
    drop_frame.pack()

    root.mainloop()

if __name__ == "__main__":
    create_gui()







